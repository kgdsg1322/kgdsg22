#!/usr/bin/env python3
"""extract_sermon.py — 설교문 결정론 추출 (LLM이 파일을 직접 해석하지 않음)

지원: .txt .md (그대로) / .docx (python-docx) / .pdf (PyMuPDF)
거부: .pages 등 그 외 (변환 요청 안내)
사용: python3 extract_sermon.py <입력파일> <출력.txt> [--min-chars 100]
출력: JSON (stdout). status=OK면 출력.txt에 전문 저장됨.
"""
import argparse, json, sys
from pathlib import Path


def extract(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="replace")
    if ext == ".docx":
        import docx
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                parts.extend(c.text for c in row.cells)
        return "\n".join(parts)
    if ext == ".pdf":
        import fitz
        doc = fitz.open(str(path))
        return "\n".join(page.get_text() for page in doc)
    raise ValueError(f"지원하지 않는 형식 '{ext}' — txt/md/docx/pdf로 변환을 요청하세요")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("input")
    ap.add_argument("output")
    ap.add_argument("--min-chars", type=int, default=100)
    a = ap.parse_args()

    src = Path(a.input).expanduser()
    if not src.exists():
        print(json.dumps({"status": "FAIL", "reason": f"파일 없음: {src}"}, ensure_ascii=False))
        return 1
    try:
        text = extract(src).strip()
    except Exception as e:
        print(json.dumps({"status": "FAIL", "reason": str(e)}, ensure_ascii=False))
        return 1
    if len(text) < a.min_chars:
        print(json.dumps({"status": "FAIL",
                          "reason": f"추출 텍스트 {len(text)}자 — 설교문으로 보기에 너무 짧음(최소 {a.min_chars}자). 임의 제작 금지 원칙에 따라 중단."},
                         ensure_ascii=False))
        return 1
    out = Path(a.output).expanduser()
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(text, encoding="utf-8")
    print(json.dumps({"status": "OK", "chars": len(text), "output": str(out)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
